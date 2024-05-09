from docx import Document
from docx.shared import Pt, Cm
# pt的作用是英镑
from docx.shared import RGBColor
from docx.oxml.ns import qn

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# 用于设置对象居中对其 \ 用于 行距调整
# import re
# import os
# path = os.getcwd()

# filename = input("请输入文件名称： ")
filename = "测试文档.docx"
doc = Document(r"C:\Users\Cc\Desktop\%s" % filename)
# for k in os.listdir(path ="C:\ Users\Cc\Desktop\%s" % filename):
#     if re.match(".+\.docx", k):
for paragraph in doc.paragraphs:
    paragraph.text = paragraph.text.lstrip()
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.text = run.text.join(run.text.split())
# 正文
for paragraph in doc.paragraphs[1:-1]:
    if len(paragraph.text) == 0:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.line_spacing = Pt(23)
    # 首行缩进
    paragraph.paragraph_format.first_line_indent = Cm(1)
    for run in paragraph.runs:
        run.font.bold = False
        run.font.italic = False
        run.font.underline = False
        run.font.strike = False
        run.font.shadow = False
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

        run.font.name = "宋体"
        # 下两行代码只针对宋体这样的中文字体
        r = run._element.rPr.rFonts
        r.set(qn("w:eastAsia"), "宋体")
# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 标题
doc.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
for run in doc.paragraphs[0].runs:
    run.font.size = Pt(18)
    run.font.bold = True
doc.paragraphs[0].paragraph_format.line_spacing = 1
doc.paragraphs[0].paragraph_format.space_before = Pt(12)
doc.paragraphs[0].paragraph_format.space_after = Pt(12)

# 落款
for run in doc.paragraphs[-1].runs:
    run.font.size = Pt(14)
    run.font.name = "仿宋"
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), "仿宋")
doc.paragraphs[-1].paragraph_format.line_spacing = 1
doc.paragraphs[-1].paragraph_format.space_before = Pt(6)
doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# 结束
doc.save(r"C:\Users\Cc\Desktop\%s" % filename)
